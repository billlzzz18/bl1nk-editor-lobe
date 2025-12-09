import os
import json
import requests
import numpy as np
from typing import List, Dict, Any, Optional
from datetime import datetime
import faiss
from sentence_transformers import SentenceTransformer
import PyPDF2
import docx
import csv
import markdown
from bs4 import BeautifulSoup
import pandas as pd
from src.models.enhanced_models import RAGDocument, UploadedFile
from src.models.user import db
from src.config import Config

class EnhancedRAGService:
    def __init__(self):
        self.embedding_model = SentenceTransformer('all-MiniLM-L6-v2')
        self.index = None
        self.documents = []
        self.notion_api_key = Config.NOTION_API_KEY
        self.initialize_faiss_index()
    
    def initialize_faiss_index(self):
        """Initialize FAISS index for vector similarity search"""
        try:
            # Try to load existing index
            if os.path.exists('faiss_index.bin'):
                self.index = faiss.read_index('faiss_index.bin')
                print("Loaded existing FAISS index")
            else:
                # Create new index (384 dimensions for all-MiniLM-L6-v2)
                self.index = faiss.IndexFlatIP(384)
                print("Created new FAISS index")
        except Exception as e:
            print(f"Error initializing FAISS index: {e}")
            self.index = faiss.IndexFlatIP(384)
    
    def save_faiss_index(self):
        """Save FAISS index to disk"""
        try:
            faiss.write_index(self.index, 'faiss_index.bin')
            print("FAISS index saved successfully")
        except Exception as e:
            print(f"Error saving FAISS index: {e}")
    
    def extract_text_from_file(self, file_path: str, file_type: str) -> str:
        """Extract text content from various file types"""
        try:
            if file_type.lower() == 'pdf':
                return self._extract_from_pdf(file_path)
            elif file_type.lower() in ['docx', 'doc']:
                return self._extract_from_docx(file_path)
            elif file_type.lower() == 'csv':
                return self._extract_from_csv(file_path)
            elif file_type.lower() == 'html':
                return self._extract_from_html(file_path)
            elif file_type.lower() == 'md':
                return self._extract_from_markdown(file_path)
            elif file_type.lower() == 'txt':
                return self._extract_from_txt(file_path)
            else:
                return ""
        except Exception as e:
            print(f"Error extracting text from {file_type} file: {e}")
            return ""
    
    def _extract_from_pdf(self, file_path: str) -> str:
        """Extract text from PDF file"""
        text = ""
        with open(file_path, 'rb') as file:
            pdf_reader = PyPDF2.PdfReader(file)
            for page in pdf_reader.pages:
                text += page.extract_text() + "\n"
        return text
    
    def _extract_from_docx(self, file_path: str) -> str:
        """Extract text from DOCX file"""
        doc = docx.Document(file_path)
        text = ""
        for paragraph in doc.paragraphs:
            text += paragraph.text + "\n"
        return text
    
    def _extract_from_csv(self, file_path: str) -> str:
        """Extract text from CSV file"""
        df = pd.read_csv(file_path)
        return df.to_string()
    
    def _extract_from_html(self, file_path: str) -> str:
        """Extract text from HTML file"""
        with open(file_path, 'r', encoding='utf-8') as file:
            soup = BeautifulSoup(file.read(), 'html.parser')
            return soup.get_text()
    
    def _extract_from_markdown(self, file_path: str) -> str:
        """Extract text from Markdown file"""
        with open(file_path, 'r', encoding='utf-8') as file:
            md_content = file.read()
            html = markdown.markdown(md_content)
            soup = BeautifulSoup(html, 'html.parser')
            return soup.get_text()
    
    def _extract_from_txt(self, file_path: str) -> str:
        """Extract text from TXT file"""
        with open(file_path, 'r', encoding='utf-8') as file:
            return file.read()
    
    def process_uploaded_file(self, file_id: int, user_id: int) -> bool:
        """Process uploaded file and add to RAG system"""
        try:
            uploaded_file = UploadedFile.query.filter_by(id=file_id, user_id=user_id).first()
            if not uploaded_file:
                return False
            
            # Extract text content
            content = self.extract_text_from_file(uploaded_file.file_path, uploaded_file.file_type)
            
            if not content.strip():
                return False
            
            # Update uploaded file with extracted content
            uploaded_file.extracted_content = content
            uploaded_file.is_processed = True
            
            # Create RAG document
            rag_doc = RAGDocument(
                user_id=user_id,
                source_type='file',
                source_id=str(file_id),
                title=uploaded_file.original_filename,
                content=content,
                metadata=json.dumps({
                    'file_type': uploaded_file.file_type,
                    'file_size': uploaded_file.file_size,
                    'original_filename': uploaded_file.original_filename
                })
            )
            
            # Generate embedding
            embedding = self.embedding_model.encode([content])[0]
            rag_doc.embedding = json.dumps(embedding.tolist())
            
            # Add to database
            db.session.add(rag_doc)
            db.session.commit()
            
            # Add to FAISS index
            self.add_document_to_index(rag_doc.id, embedding, content)
            
            return True
            
        except Exception as e:
            print(f"Error processing uploaded file: {e}")
            db.session.rollback()
            return False
    
    def add_document_to_index(self, doc_id: int, embedding: np.ndarray, content: str):
        """Add document to FAISS index"""
        try:
            # Ensure embedding is 2D array
            if embedding.ndim == 1:
                embedding = embedding.reshape(1, -1)
            
            # Add to FAISS index
            self.index.add(embedding.astype('float32'))
            
            # Store document info
            self.documents.append({
                'id': doc_id,
                'content': content[:500] + '...' if len(content) > 500 else content
            })
            
            # Save index
            self.save_faiss_index()
            
        except Exception as e:
            print(f"Error adding document to index: {e}")
    
    def search_notion_pages(self, query: str, user_id: int) -> List[Dict[str, Any]]:
        """Search Notion pages using Notion API"""
        if not self.notion_api_key:
            return []
        
        try:
            headers = {
                'Authorization': f'Bearer {self.notion_api_key}',
                'Content-Type': 'application/json',
                'Notion-Version': '2022-06-28'
            }
            
            # Search for pages
            search_data = {
                'query': query,
                'filter': {
                    'value': 'page',
                    'property': 'object'
                }
            }
            
            response = requests.post(
                'https://api.notion.com/v1/search',
                headers=headers,
                json=search_data
            )
            
            if response.status_code == 200:
                results = response.json().get('results', [])
                pages = []
                
                for result in results:
                    page_info = {
                        'id': result.get('id'),
                        'title': self._extract_notion_title(result),
                        'url': result.get('url'),
                        'last_edited': result.get('last_edited_time'),
                        'created': result.get('created_time')
                    }
                    pages.append(page_info)
                
                return pages
            else:
                print(f"Notion API error: {response.status_code}")
                return []
                
        except Exception as e:
            print(f"Error searching Notion pages: {e}")
            return []
    
    def _extract_notion_title(self, page_data: Dict[str, Any]) -> str:
        """Extract title from Notion page data"""
        try:
            properties = page_data.get('properties', {})
            
            # Try to find title property
            for prop_name, prop_data in properties.items():
                if prop_data.get('type') == 'title':
                    title_array = prop_data.get('title', [])
                    if title_array:
                        return title_array[0].get('text', {}).get('content', 'Untitled')
            
            # Fallback to page title
            return page_data.get('title', [{}])[0].get('text', {}).get('content', 'Untitled')
            
        except Exception as e:
            print(f"Error extracting Notion title: {e}")
            return 'Untitled'
    
    def get_notion_page_content(self, page_id: str) -> str:
        """Get content from a specific Notion page"""
        if not self.notion_api_key:
            return ""
        
        try:
            headers = {
                'Authorization': f'Bearer {self.notion_api_key}',
                'Content-Type': 'application/json',
                'Notion-Version': '2022-06-28'
            }
            
            # Get page blocks
            response = requests.get(
                f'https://api.notion.com/v1/blocks/{page_id}/children',
                headers=headers
            )
            
            if response.status_code == 200:
                blocks = response.json().get('results', [])
                content = ""
                
                for block in blocks:
                    block_content = self._extract_block_content(block)
                    if block_content:
                        content += block_content + "\n"
                
                return content
            else:
                print(f"Error getting Notion page content: {response.status_code}")
                return ""
                
        except Exception as e:
            print(f"Error getting Notion page content: {e}")
            return ""
    
    def _extract_block_content(self, block: Dict[str, Any]) -> str:
        """Extract text content from Notion block"""
        try:
            block_type = block.get('type')
            
            if block_type in ['paragraph', 'heading_1', 'heading_2', 'heading_3']:
                rich_text = block.get(block_type, {}).get('rich_text', [])
                return ''.join([text.get('text', {}).get('content', '') for text in rich_text])
            
            elif block_type == 'bulleted_list_item':
                rich_text = block.get('bulleted_list_item', {}).get('rich_text', [])
                return '• ' + ''.join([text.get('text', {}).get('content', '') for text in rich_text])
            
            elif block_type == 'numbered_list_item':
                rich_text = block.get('numbered_list_item', {}).get('rich_text', [])
                return '1. ' + ''.join([text.get('text', {}).get('content', '') for text in rich_text])
            
            return ""
            
        except Exception as e:
            print(f"Error extracting block content: {e}")
            return ""
    
    def add_notion_page_to_rag(self, page_id: str, user_id: int) -> bool:
        """Add Notion page to RAG system"""
        try:
            # Get page content
            content = self.get_notion_page_content(page_id)
            
            if not content.strip():
                return False
            
            # Check if already exists
            existing = RAGDocument.query.filter_by(
                user_id=user_id,
                source_type='notion',
                source_id=page_id
            ).first()
            
            if existing:
                # Update existing document
                existing.content = content
                existing.updated_at = datetime.utcnow()
                
                # Update embedding
                embedding = self.embedding_model.encode([content])[0]
                existing.embedding = json.dumps(embedding.tolist())
            else:
                # Create new document
                rag_doc = RAGDocument(
                    user_id=user_id,
                    source_type='notion',
                    source_id=page_id,
                    title=f"Notion Page {page_id}",
                    content=content,
                    metadata=json.dumps({'page_id': page_id})
                )
                
                # Generate embedding
                embedding = self.embedding_model.encode([content])[0]
                rag_doc.embedding = json.dumps(embedding.tolist())
                
                db.session.add(rag_doc)
            
            db.session.commit()
            
            # Add to FAISS index
            self.add_document_to_index(
                existing.id if existing else rag_doc.id,
                embedding,
                content
            )
            
            return True
            
        except Exception as e:
            print(f"Error adding Notion page to RAG: {e}")
            db.session.rollback()
            return False
    
    def semantic_search(self, query: str, user_id: int, top_k: int = 5) -> List[Dict[str, Any]]:
        """Perform semantic search using RAG system"""
        try:
            # Generate query embedding
            query_embedding = self.embedding_model.encode([query])[0]
            
            # Search in FAISS index
            if self.index.ntotal == 0:
                return []
            
            # Ensure query embedding is 2D
            if query_embedding.ndim == 1:
                query_embedding = query_embedding.reshape(1, -1)
            
            scores, indices = self.index.search(query_embedding.astype('float32'), top_k)
            
            # Get documents from database
            results = []
            for i, (score, idx) in enumerate(zip(scores[0], indices[0])):
                if idx < len(self.documents):
                    doc_info = self.documents[idx]
                    
                    # Get full document from database
                    rag_doc = RAGDocument.query.filter_by(
                        id=doc_info['id'],
                        user_id=user_id
                    ).first()
                    
                    if rag_doc:
                        result = rag_doc.to_dict()
                        result['similarity_score'] = float(score)
                        results.append(result)
            
            return results
            
        except Exception as e:
            print(f"Error in semantic search: {e}")
            return []
    
    def get_rag_context(self, query: str, user_id: int, max_context_length: int = 2000) -> str:
        """Get relevant context for RAG-enhanced chat"""
        try:
            # Search for relevant documents
            search_results = self.semantic_search(query, user_id, top_k=3)
            
            context = ""
            current_length = 0
            
            for result in search_results:
                content = result.get('content', '')
                
                # Add content if it fits within max length
                if current_length + len(content) <= max_context_length:
                    context += f"Document: {result.get('title', 'Unknown')}\n"
                    context += f"Content: {content}\n\n"
                    current_length += len(content)
                else:
                    # Add partial content
                    remaining_length = max_context_length - current_length
                    if remaining_length > 100:  # Only add if meaningful length remains
                        context += f"Document: {result.get('title', 'Unknown')}\n"
                        context += f"Content: {content[:remaining_length]}...\n\n"
                    break
            
            return context
            
        except Exception as e:
            print(f"Error getting RAG context: {e}")
            return ""
    
    def rebuild_index(self, user_id: Optional[int] = None):
        """Rebuild FAISS index from database"""
        try:
            # Query documents
            query = RAGDocument.query
            if user_id:
                query = query.filter_by(user_id=user_id)
            
            documents = query.all()
            
            # Create new index
            self.index = faiss.IndexFlatIP(384)
            self.documents = []
            
            # Add all documents to index
            for doc in documents:
                if doc.embedding:
                    embedding = np.array(json.loads(doc.embedding))
                    self.add_document_to_index(doc.id, embedding, doc.content)
            
            print(f"Rebuilt FAISS index with {len(documents)} documents")
            
        except Exception as e:
            print(f"Error rebuilding index: {e}")

# Global instance
enhanced_rag_service = EnhancedRAGService()

